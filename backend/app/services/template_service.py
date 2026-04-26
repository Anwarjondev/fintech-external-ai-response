import os
import subprocess
import uuid
from datetime import datetime
from docx import Document

TEMPLATE_PATH = "/home/anvarjon/fintech-hackathon/fintech-external-ai-response/backend/bank_response_template-2.docx"
OUTPUT_DIR = "/home/anvarjon/fintech-hackathon/fintech-external-ai-response/backend/generated_docs"

os.makedirs(OUTPUT_DIR, exist_ok=True)

def fill_bank_template(data: dict, convert_to_pdf: bool = False) -> str:
    if not os.path.exists(TEMPLATE_PATH):
        # Fallback if template is missing
        doc = Document()
        doc.add_heading("Bank Javobi", 0)
        doc.add_paragraph(f"Mavzu: {data.get('topic')}")
        doc.add_paragraph(f"Javob: {data.get('response_text')}")
    else:
        doc = Document(TEMPLATE_PATH)
    
    # 1. Fill Mavzu and Organization
    for p in doc.paragraphs:
        if "MAVZU:" in p.text:
            p.text = f"MAVZU: {data.get('topic', 'Noma\'lum')}"
        if "Tashkilot:" in p.text:
            p.text = f"Tashkilot: {data.get('organization', 'Noma\'lum')}"
            
    # 2. Fill Hurmatli
    for p in doc.paragraphs:
        if "Hurmatli" in p.text:
            name = data.get('user_name') or "Mijoz"
            p.text = f"Hurmatli {name},"
            break
            
    # 3. Fill Date and No
    now = datetime.now()
    uz_months = {
        1: "yanvar", 2: "fevral", 3: "mart", 4: "aprel",
        5: "may", 6: "iyun", 7: "iyul", 8: "avgust",
        9: "sentyabr", 10: "oktyabr", 11: "noyabr", 12: "dekabr"
    }
    
    def fill_date_in_text(text):
        if "Sana:" in text:
            return f"Sana: «{now.day:02d}» {uz_months[now.month]} {now.year}-y."
        if "Tashkilot:" in text:
            return f"Tashkilot: {data.get('organization', 'Noma\'lum')}"
        if "yildagi №" in text:
            return f"Sizning «{now.day:02d}» {uz_months[now.month]} {now.year}-yildagi № {data.get('request_id', '000')} raqamli murojaatingizga javoban ma'lumot beramiz:"
        return text

    # Search in paragraphs
    for p in doc.paragraphs:
        new_text = fill_date_in_text(p.text)
        if new_text != p.text:
            p.text = new_text

    # Search in tables (Sana might be in a header table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    new_text = fill_date_in_text(p.text)
                    if new_text != p.text:
                        p.text = new_text
            
    # 4. Fill Answer
    answer_text = data.get('response_text', '')
    answer_started = False
    answer_paragraph = None
    for p in doc.paragraphs:
        if "JAVOB MATNI:" in p.text:
            answer_started = True
            continue
        if answer_started and "HUQUQIY ASOS:" in p.text:
            answer_started = False
            continue
        if answer_started and "___" in p.text:
            if answer_text:
                p.text = answer_text
                answer_paragraph = p
                answer_text = "" # Put only once
            else:
                p.text = ""
                
    # 4a. Add Table if data found
    retrieved = data.get("retrieved_data")
    if retrieved and retrieved.get("found") and retrieved.get("data"):
        if answer_paragraph:
            # Add a spacer paragraph after the answer
            p_msg = answer_paragraph.insert_paragraph_before("\nTopilgan ma'lumotlar jadvali:")
            answer_paragraph._p.addnext(p_msg._p) # Move it after
            
            items = retrieved["data"]
            if items and isinstance(items, list):
                keys = list(items[0].keys())
                # Create table at the end and then move it
                table = doc.add_table(rows=1, cols=len(keys))
                table.style = 'Table Grid'
                
                # Move table after p_msg
                p_msg._p.addnext(table._tbl)
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(keys):
                    hdr_cells[i].text = str(key).replace("_", " ").title()
                
                # Rows
                for item in items:
                    row_cells = table.add_row().cells
                    for i, key in enumerate(keys):
                        val = item.get(key)
                        if isinstance(val, (int, float)):
                            row_cells[i].text = f"{val:,.2f}"
                        else:
                            row_cells[i].text = str(val)
                
    # 5. Fill Legal Basis
    legal_ref = data.get('legal_reference', '')
    legal_started = False
    for p in doc.paragraphs:
        if "HUQUQIY ASOS:" in p.text:
            legal_started = True
            continue
        if legal_started and "Yuqoridagilarni inobatga" in p.text:
            legal_started = False
            continue
        if legal_started and "___" in p.text:
            if legal_ref:
                p.text = legal_ref
                legal_ref = ""
            else:
                p.text = ""
                
    # 6. Footer
    for p in doc.paragraphs:
        if "Hujjat ID:" in p.text:
            p.text = p.text.replace("___________", str(data.get('request_id', '000')), 1)
            p.text = p.text.replace("___________", now.strftime("%Y%m%d%H%M"), 1)
            
    # Save DOCX
    filename = f"response_{data.get('request_id', 'req')}_{uuid.uuid4().hex[:6]}"
    docx_path = os.path.join(OUTPUT_DIR, f"{filename}.docx")
    doc.save(docx_path)
    
    if convert_to_pdf:
        try:
            # LibreOffice conversion
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf", 
                "--outdir", OUTPUT_DIR, docx_path
            ], check=True, capture_output=True)
            return os.path.join(OUTPUT_DIR, f"{filename}.pdf")
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            return docx_path
            
    return docx_path
